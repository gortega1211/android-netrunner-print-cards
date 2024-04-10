from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.shared import Cm
from io import BytesIO
import os
import re
import requests

# https://www.crummy.com/software/BeautifulSoup/bs4/doc/#searching-the-tree

CARD_WIDTH = 6.2
CARD_HEIGHT = 9.0

PATH = "./documents"
TEMPLATE_FILENAME = "template.docx"

NETRUNNER_SETS = {
    "core2": "/en/set/core2", # Revised Core Set
    "kitara": "/en/cycle/kitara", # Kitara
    "rar": "/en/set/rar", # Reign and Reverie
    "mo": "/en/set/mo", # Magnus Opus
    "sc19": "/en/set/sc19", # System Core 2019
    "mor": "/en/set/mor", # Magnus Opus Reprint
    "sm": "/en/set/sm", # Salvaged Memories
    "sg": "/en/set/sg", # System Gateway
    "su21": "/en/set/su21", # System Update 2021
    "ur": "/en/set/ur", # Ashes | Uprising
    "urbp": "/en/set/urbp", # Ashes | Uprising Booster Pack
    "df": "/en/set/df", # Ashes | Downfall
    "ph": "/en/set/ph", # Borealis | Parhelion
    "ms": "/en/set/ms", # Borealis | Midnight Sun
    "msbp": "/en/set/msbp", # Borealis | Midnight Sun Booster Pack
    "tai": "/en/set/tai", # Liberation | The Automata Initiative
}

OPTIONS_CARD = [
    "acoo_promos",
    "nisei"
]

BEST_RESOLUTIONS_LINKS = (
    "https://www.jinteki.net/img/cards/en/high/stock/",
    "https://www.jinteki.net/img/cards/en/default/stock/"
)

IMG_EXTENSIONS = (
    ".png",
    ".jpg",
)

DOUBLE_TYPES = {
    "-front",
    "-back",
}

def _verify_url(url):
    resp = requests.get(url)
    return resp.status_code == 200

def get_best_image_resolution(id_card):
    for brl in BEST_RESOLUTIONS_LINKS:
        url = brl + id_card
        for ie in IMG_EXTENSIONS:
            if _verify_url(url + ie):
                return url + ie
    return False

def get_images(option_card, url_base, soup):
    items = []
    # match
    # case
    if option_card == "acoo_promos":
        for i in soup.find_all(style=re.compile("width")):
            print(i["alt"])
            items.append({
                "image": url_base + i["src"]
            })
    elif option_card == "nisei":
        for i in soup.find_all("a", {"class": "card"}):
            id_card = i.get("href").split("/")[-1]
            soup = get_page_soup(i.get("href"))
            img_class = soup.find("img", {"class": "card-image"})
            print(img_class["alt"])
            img_url = img_class["data-src"]
            card_text = soup.find("div", {"class": "card-text"})
            list_urls = []
            if card_text and 'flip this' in card_text.p.text:
                for dt in DOUBLE_TYPES:
                    list_urls.append(get_best_image_resolution(id_card + dt))
            else:
                list_urls.append(get_best_image_resolution(id_card))
            for br_img_url in list_urls:
                for _ in range(3):
                    items.append({
                        "image": br_img_url if br_img_url else img_url
                    })
    return items

def get_page_soup(url):
    resp = requests.get(url)
    if resp.status_code == 404:
        raise ValueError(f"The page '{url}' was not found.")
    content = resp.content
    soup = BeautifulSoup(content, 'lxml')
    return soup

def request_input(text):
    user_input = input(text)
    return user_input

def copy_docx_template(dest):
    import shutil

    src = os.path.join(PATH, TEMPLATE_FILENAME)
    shutil.copyfile(src, dest)

def put_images_in_word(path, data):
    for idx, i in enumerate(data):
        if not idx % 9:
            with open(path, 'rb') as f:
                document = Document(f)
            p = document.add_paragraph()
            r = p.add_run()
        image_from_url = requests.get(i["image"])
        binary_img = BytesIO(image_from_url.content)
        r.add_picture(binary_img, width=Cm(CARD_WIDTH), height=Cm(CARD_HEIGHT))
        r.add_text(" ")
        if idx and (not (idx + 1) % 9 or len(data)-1 == idx):
            document.save(path)

def run():
    option_card = request_input("What cards do you want to get?:\n* " + '\n* '.join(OPTIONS_CARD) + "\n-> ")
    if option_card == "acoo_promos":
        url_base = "https://acoo.net"
        url = url_base + "/netrunner/cards/set/promos/"
    elif option_card == "nisei":
        url_base = "https://netrunnerdb.com"
        url = url_base + NETRUNNER_SETS["core2"] # <-- Ingresar Key (Ver arriba)
        
    page_soup = get_page_soup(url)
    data = get_images(option_card, url_base, page_soup)
    filename = request_input("What is the filename?: ")
    path = os.path.join(PATH, filename + "_cards.docx")
    copy_docx_template(path)
    put_images_in_word(path, data)

if __name__ == "__main__":
    run()